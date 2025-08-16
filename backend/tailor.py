# backend/tailor.py
from __future__ import annotations
import os, json, re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Tuple
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from skills_taxonomy import detect_role

# Paths
DATA_DIR = Path("data")  # served by FastAPI via /files
TEMPLATES_DIR = (Path(__file__).resolve().parent / "templates" / "resume")
DATA_DIR.mkdir(parents=True, exist_ok=True)

def _nowstamp() -> str:
    return datetime.utcnow().strftime("%Y%m%d-%H%M%S")

def _slug(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", (s or "doc").lower()).strip("-")

def _read_template(role: str) -> Dict[str, Any]:
    # Prefer on-disk template: backend/templates/resume/<role>.json
    path = TEMPLATES_DIR / f"{(role or 'devops').lower()}.json"
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))

    # Fallback to devops.json if specific role missing
    fallback = TEMPLATES_DIR / "devops.json"
    if fallback.exists():
        return json.loads(fallback.read_text(encoding="utf-8"))

    # Final minimal inline fallback so Tailor never 500s
    return {
        "summary": "Results-driven engineer with experience aligned to the role.",
        "core_skills": ["AWS", "Kubernetes", "Terraform", "CI/CD", "Linux"],
        "bullets": [
            {"text": "Implemented CI/CD pipelines improving deployment frequency by {X}%.", "tags": ["ci", "cd"]},
            {"text": "Managed Kubernetes clusters and IaC with Terraform across {N} environments.", "tags": ["kubernetes", "terraform"]},
            {"text": "Built monitoring with Prometheus/Grafana reducing MTTR by {X}%.", "tags": ["prometheus", "grafana"]},
        ],
    }

def _tokenize(text: str) -> List[str]:
    return re.findall(r"[a-zA-Z0-9\+#\.]+", (text or "").lower())

def _extract_keywords(jd: str) -> List[str]:
    toks = _tokenize(jd)
    from collections import Counter
    cnt = Counter(toks)
    return [w for w, n in cnt.most_common() if len(w) > 2][:200]

def _choose_bullets(template: Dict[str, Any], jd_keywords: List[str], limit: int = 6) -> List[str]:
    jd_set = set(jd_keywords)
    scored: List[Tuple[int, str]] = []
    for b in template.get("bullets", []):
        tags = set(t.lower() for t in b.get("tags", []))
        overlap = len(tags & jd_set)
        scored.append((overlap, b["text"]))
    scored.sort(key=lambda x: x[0], reverse=True)

    out: List[str] = []
    for i, (_, text) in enumerate(scored):
        if i >= limit:
            break
        # simple placeholder fills so output is readable without extra data
        text = (
            text.replace("{X}", "30")
                .replace("{N}", "5")
                .replace("{M}", "60")
                .replace("{period}", "last year")
        )
        out.append(text)
    return out

def _profile_from_env(profile: Dict[str, Any] | None) -> Dict[str, Any]:
    p = dict(profile or {})
    p.setdefault("first_name", os.getenv("PROFILE_FIRST_NAME", ""))
    p.setdefault("last_name",  os.getenv("PROFILE_LAST_NAME", ""))
    p.setdefault("email",      os.getenv("PROFILE_EMAIL", ""))
    p.setdefault("phone",      os.getenv("PROFILE_PHONE", ""))
    p.setdefault("linkedin",   os.getenv("PROFILE_LINKEDIN", ""))
    p.setdefault("resume_path",os.getenv("PROFILE_RESUME", ""))
    p.setdefault("role",       os.getenv("PROFILE_ROLE", ""))
    return p

def _make_ats_docx_resume(
    out_path: Path,
    job: Dict[str, Any],
    profile: Dict[str, Any],
    template: Dict[str, Any],
    picked_bullets: List[str],
    jd_keywords: List[str],
):
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Header
    full_name = f"{profile.get('first_name','')} {profile.get('last_name','')}".strip()
    h = doc.add_paragraph()
    run = h.add_run(full_name or "Your Name")
    run.bold = True
    run.font.size = Pt(16)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.add_run(f"{profile.get('email','')} • {profile.get('phone','')} • {profile.get('linkedin','')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    # Summary
    doc.add_paragraph().add_run("Summary").bold = True
    doc.add_paragraph(template.get("summary", ""))

    # Core Skills (ATS-friendly)
    doc.add_paragraph().add_run("Core Skills").bold = True
    skills = template.get("core_skills", []) or []
    addl = [k for k in jd_keywords if k.upper() not in [s.upper() for s in skills]][:8]
    skills_line = ", ".join(skills + addl)
    doc.add_paragraph(skills_line)

    # Experience Highlights
    doc.add_paragraph().add_run("Experience Highlights").bold = True
    for b in picked_bullets:
        doc.add_paragraph(f"• {b}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

def _cover_letter_text(
    company: str,
    title: str,
    template: Dict[str, Any],
    picked_bullets: List[str],
    profile: Dict[str, Any],
) -> str:
    full_name = f"{profile.get('first_name','')} {profile.get('last_name','')}".strip() or "Your Name"
    lines = [
        f"Dear Hiring Team at {company or 'the company'},",
        "",
        f"I’m excited to apply for the {title or 'role'} position. My background aligns closely with your needs, particularly around the areas highlighted in the job description.",
        "",
        "A few relevant highlights:",
    ]
    for b in picked_bullets[:3]:
        lines.append(f"• {b}")
    lines += [
        "",
        "I value pragmatic automation, clear SLOs, and collaborative delivery with security as a first principle. I’d welcome the chance to talk through how I can help your team achieve its goals.",
        "",
        f"Best regards,\n{full_name}\n{profile.get('email','')} | {profile.get('phone','')}",
    ]
    return "\n".join(lines)

def _ats_score(core_skills: List[str], picked_bullets: List[str], jd_keywords: List[str]) -> int:
    jd = set(k.lower() for k in jd_keywords)
    skills = set((core_skills or []) and [s.lower() for s in core_skills])
    # skill coverage
    skill_cov = len(jd & skills)
    # bullet keyword coverage
    bullet_terms: set[str] = set()
    for b in picked_bullets:
        for t in re.findall(r"[a-zA-Z0-9\+#\.]+", b.lower()):
            if len(t) > 2:
                bullet_terms.add(t)
    bullet_cov = len(jd & bullet_terms)
    # heuristic: weight skills higher
    raw = (2 * skill_cov) + bullet_cov
    # normalize roughly against jd size
    denom = max(8, len(jd) // 4)
    return max(0, min(100, int((raw / denom) * 100)))

def tailor(job: Dict[str, Any], profile: Dict[str, Any] | None = None) -> Dict[str, Any]:
    """
    Build ATS-friendly resume + cover letter tailored to a single job.
    Returns file paths under data/ (your API exposes them as /files/...).
    """
    profile = _profile_from_env(profile)
    title = (job.get("title") or "").strip()
    company = (job.get("company") or "").strip()
    jd = (job.get("jd_text") or "").strip()

    # 1) detect role & load template
    role = detect_role(title, jd, explicit_role=profile.get("role"))
    tpl = _read_template(role)

    # 2) parse JD keywords & pick bullets
    jd_keywords = _extract_keywords(jd)
    picked_bullets = _choose_bullets(tpl, jd_keywords, limit=6)

    # 3) generate ATS resume
    stamp = _nowstamp()
    base = f"{_slug(company)}-{_slug(title)}-{stamp}"
    resume_path = DATA_DIR / "resumes" / f"{base}.docx"
    _make_ats_docx_resume(resume_path, job, profile, tpl, picked_bullets, jd_keywords)

    # 4) generate cover letter (txt + docx)
    cl_text = _cover_letter_text(company, title, tpl, picked_bullets, profile)
    cl_txt_path = DATA_DIR / "cover_letters" / f"{base}.txt"
    cl_txt_path.parent.mkdir(parents=True, exist_ok=True)
    cl_txt_path.write_text(cl_text, encoding="utf-8")

    cl_docx_path = DATA_DIR / "cover_letters" / f"{base}.docx"
    cl_doc = Document()
    for line in cl_text.split("\n"):
        cl_doc.add_paragraph(line)
    cl_doc.save(str(cl_docx_path))

    # 5) score for UI
    ats = _ats_score(tpl.get("core_skills", []), picked_bullets, jd_keywords)

    return {
        "role_detected": role,
        "revised_bullets": picked_bullets,
        "summary": tpl.get("summary"),
        "core_skills": tpl.get("core_skills"),
        "resume_docx_path": str(resume_path),
        "cover_letter_path": str(cl_docx_path),      # docx (upload-friendly)
        "cover_letter_text_path": str(cl_txt_path),  # plaintext preview
        "ats_score": ats,                             # NEW
    }
